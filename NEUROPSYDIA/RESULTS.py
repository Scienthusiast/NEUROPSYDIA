# -*- coding: cp1252 -*-
import os,sys
sys.path.append('.\\Files\\Docs')
import neuropsydia as n
from neuropsydia import *
from docx import Document
from docx.shared import Inches

origin=os.getcwd()


#INFOS
version = 'v0.1'
authors = 'D. Makowski, L. Dutriaux et al. (2016)'
autorisations = n.autorisations







def identification(civilite="M."):

    boucle=True
    while boucle == True:
        for event in pygame.event.get():
            if event.type == KEYDOWN and event.key == K_ESCAPE:
            	boucle = False

        
        
        screen.fill((255, 255, 255))

        ecrire('Identification','arialblack',W/20,(W/2, H/9*1),(0,0,0))
        Neuropsychologist = text_input('Mot de passe ? : ',(W/5,H/9*2),allowable=autorisations)









        Test = text_input('Test ? (o/n) : ',(W/8.5,H/9*2.75),(0,0,0))
        if Test == 'n':
            Initials = text_input('Initiales : ',(W/8.5,H/9*3.5),(0,0,0))
            Sex = text_input('Sexe (m/f) : ',(W/8.5,H/9*4.5),(0,0,0))
            Birth_day = text_input('Jour de naissance (jj) : ',(W/8.5,H/9*5.5),(0,0,0))
            Birth_month = text_input('Mois de naissance (mm) : ',(W/8.5,H/9*6.5),(0,0,0))
            Birth_year = text_input('Année de naissance (aaaa) : ',(W/8.5,H/9*7.5),(0,0,0))
            subject_code = str(Birth_year) + '_' + str(Birth_month) + '_' + str(Birth_day) + '_' + Initials + '_' + Sex
        else:
            subject_code= "XXXX_XX_XX_XX_x"
            Sex = 'm'
            Initials = 'XX'
        if os.path.isfile(os.getcwd().split('\\Files')[0] + '/Data/'+subject_code+'/'+subject_code+'_INFOS.csv') == False:
            ecrire('Sujet inconnu. Appuyez sur ENTRER pour recommencer.','arialblack',W/40,(W/2, H/9*8.5),(255,0,0),passer="K_RETURN")
            identification()



        if Sex == 'f': civilite = 'Mme.'


        return(subject_code,civilite,Initials)


def launcher((column,row),test=[],scale = W/7):

    pos_images=[]
    number=0
    screen.fill((255, 255, 255))
    for j in range(row):
        for i in range(column):
            if (number+1)>len(test):
                pass
            else:
                limage=pygame.image.load('./Files/Docs/Logo_'+test[number]+'.png')
                large=limage.get_width()
                haute=limage.get_height()
                limage=pygame.transform.smoothscale(limage,(large*scale/H,haute*scale/H))
                rectangle = limage.get_rect()
                rectangle.center = (((W/(column+1))*(i+1)),((H/(row+1))*(j+1)))
                pos_images.append((((W/(column+1))*(i+1)),((H/(row+1))*(j+1))))
                screen.blit(limage,rectangle)
                number=number+1
    pygame.display.flip()



#==============================================================================
# TRY
#==============================================================================
try:
    pygame.init()
    screen.fill((255, 255, 255))
    affiche('Logo_Launcher',(W/2,H/2.40),scale=W/12,path='./Files/Docs/')
    ecrire('Appuyez sur ENTREE pour commencer.','arialblack',W/50,(W/2, H/9*8.5),(0,0,0),passer="K_RETURN")
    pygame.display.flip()






    subject_code,civilite,Initials = identification()



    tests = []
    if os.path.isfile(os.getcwd().split('\\Files')[0] + '/Data/'+subject_code+'/'+subject_code+'_BPDQ.csv') == True: tests.append("BPDQ")
    if os.path.isfile(os.getcwd().split('\\Files')[0] + '/Data/'+subject_code+'/'+subject_code+'_Switch.csv') == True: tests.append("Switch")
    if os.path.isfile(os.getcwd().split('\\Files')[0] + '/Data/'+subject_code+'/'+subject_code+'_SimAct.csv') == True: tests.append("SimAct")
    if os.path.isfile(os.getcwd().split('\\Files')[0] + '/Data/'+subject_code+'/'+subject_code+'_Wordini.csv') == True: tests.append("Wordini")
    if os.path.isfile(os.getcwd().split('\\Files')[0] + '/Data/'+subject_code+'/'+subject_code+'_RIDE.csv') == True: tests.append("RIDE")
    if os.path.isfile(os.getcwd().split('\\Files')[0] + '/Data/'+subject_code+'/'+subject_code+'_IPIP6.csv') == True: tests.append("IPIP6")


    screen.fill((255, 255, 255))
    launcher((2,3),test=tests,scale = W/7)
    ecrire(civilite + ' ' + Initials + ' a réalisé(e) les ' + str(len(tests)) + ' épreuves suivantes :','arialblack',W/50,(W/2, H/9*1),(0,0,0))
    ecrire("Appuyez sur ENTRER pour continuer.",'arialblack',W/50,(W/2, H/9*8),(0,0,0),passer="K_RETURN")


#==============================================================================
# DATABASE Selection
#==============================================================================

    df = pandas.read_csv(os.getcwd().split('\\Files')[0] + '/Data/'+subject_code+'/'+subject_code+'_SUMMARY.csv',header = 0,engine='python',sep=';',decimal='.',index_col="subject_code")
    db = pandas.read_csv(os.getcwd().split('\\Files')[0] + '/Data/Database.csv',header = 0,engine='python',sep=';',decimal='.')
    
    screen.fill((255, 255, 255))
    ecrire("Souhaitez-vous comparer le sujet par rapport à",'arialblack',W/50,(W/2, H/9*2),(0,0,0))
    general_db = text_input('la population générale (o/n) : ',(W/8.5,H/9*3),(0,0,0),allowable=['o','n'])
    if general_db == 'n':
        control_age = text_input("Contrôler l'âge ? (o/n) : ",(W/8.5,H/9*4),(0,0,0),allowable=['o','n'])
        control_sex = text_input("Contrôler le sexe ? (o/n) : ",(W/8.5,H/9*5),(0,0,0),allowable=['o','n'])
#        control_sl = text_input("Contrôler le niveau d'étude ? (o/n) : ",(W/8.5,H/9*6),(0,0,0),allowable=['o','n'])
    
        if control_age == 'o': db = db[(db.Age>=df.Age[0] - 5) & (db.Age<=df.Age[0] + 5)]
        if control_sex == 'o': db = db[(db.Sex==df.Sex[0])]
#        if control_sl == 'o': db = db[(db.Current_study_level == df.Current_study_level[0])]
            
    








    
    



#==============================================================================
# Neuropsychologue
#==============================================================================
#    if df.Neuropsychologist[0] == 'dom':
#        df.Neuropsychologist[0] = u'D. Makowski, Neuropsychologue Certifié'
#    elif df.Neuropsychologist[0] == 'leo':
#        df.Neuropsychologist[0] = u'le dr. L. Dutriaux, Chercheur en sciences cognitives et gros ouf de la life'
#    else:
#        df.Neuropsychologist[0] = u'un utilisateur non enregistré'
    if autorisations.has_key(df.Neuropsychologist[0])==True:df.Neuropsychologist[0] = autorisations[df.Neuropsychologist[0]]
    else: df.Neuropsychologist[0] = u'un utilisateur non enregistré'

#==============================================================================
# #AGE
#==============================================================================

    ggsave(ggplot(db, aes(x='Age')) + \
    geom_density(alpha=0.3, fill="#3399FF") +\
    geom_vline(xintercept=df.Age,color="#FF0066",alpha=20) +\
    theme_bw(),"./Files/Docs/Temporary/Age.png")
    
    
    screen.fill((255, 255, 255))
    ecrire('Age','arialblack',W/30,(W/2, H/9*1),(0,0,0))

    affiche('Age',(W/2,H/9*6),scale=W/9*1,path='./Files/Docs/Temporary/')
    ecrire('Comparé à ' + str(len(db.Age)) + ' individus, ' + civilite + ' ' + Initials + ', agé(e) de ' + str(int(np.average(df.Age))) + ' ans,','arialblack',W/70,(W/2, H/9*2.2),(0,0,0))
    ecrire('se situe à ' + str("{0:.2f}".format((np.average(df.Age)-np.average(db.Age))/np.std(db.Age))) + ' ety de la moyenne du groupe (M = '+str("{0:.2f}".format(np.average(db.Age)))+' ans).','arialblack',W/70,(W/2, H/9*2.6),(0,0,0),passer="K_RETURN")




#==============================================================================
# #IPIP6
#==============================================================================

#http://docs.ggplot2.org/current/coord_polar.html
#
#    ggsave(ggplot(df, aes(x=df.IPIP6_Extraversion[0])) +\
#    geom_bar() +\
#    theme_bw(),"./Files/Docs/Temporary/Personality.png")



#==============================================================================
# COMPTE RENDU
#==============================================================================
    screen.fill((255, 255, 255))
    CR = text_input('Créer un compte-rendu ? (o/n) : ',(W/8.5,H/9*3),(0,0,0),allowable=['o','n'])
    pygame.display.flip()
    
    if CR == 'o':
        CR = Document()
        CR.add_heading('Examen Neuropsychologique de ' + civilite + ' ' + Initials, 0)
        CR.add_paragraph(civilite + ' ' + Initials + u", agé(e) de " + str(int(np.average(df.Age))) +  ' ans,'\
        + u' a été(e) évalué(e) le ' + str(df.Examination_date[0]) + u' à ' + str(df.Examination_time[0]) + ' par ' \
        + df.Neuropsychologist[0] + '.')
        
        CR.add_heading(u'Informations générales', level=1)
        CR.add_paragraph(civilite + ' ' + Initials + u', agé(e) de ' + str(int(np.average(df.Age))) + u' ans (né(e) le' \
        + ' ' + str(df.Birth_day[0]) + '/'+ str(df.Birth_month[0]) +'/'+ str(df.Birth_year[0]) + u').')



        CR.add_heading(u'Antécédents personnels et familiaux', level=1)
        if pandas.isnull(df.Family_history_disorder_1[0]) == False:
            CR.add_paragraph(u'Antécédents familiaux: ' + str(df.Family_history_disorder_1[0]) + u" (degré de certitude: " + \
            str(df.Family_history_confidence_1[0]) + u"/3), chez un membre au " + str(df.Family_history_degree_1[0]) + u' degré coté ' +\
            str(df.Family_history_side_1[0]) + u'.')
        else: CR.add_paragraph(u"Absence d'antécédents familiaux.")
        if pandas.isnull(df.Family_history_disorder_2[0]) == False:
            CR.add_paragraph(u'Autre antécédents familiaux: ' + str(df.Family_history_disorder_2[0]) + u" (degré de certitude: " + \
            str(df.Family_history_confidence_2[0]) + u"/3), chez un membre au " + str(df.Family_history_degree_2[0]) + u' degré coté ' +\
            str(df.Family_history_side_2[0]) + u'.')
        if pandas.isnull(df.Family_history_disorder_3[0]) == False:
            CR.add_paragraph(u'Autre antécédents familiaux: ' + str(df.Family_history_disorder_3[0]) + u" (degré de certitude: " + \
            str(df.Family_history_confidence_3[0]) + u"/3), chez un membre au " + str(df.Family_history_degree_3[0]) + u' degré coté ' +\
            str(df.Family_history_side_3[0]) + u'.')
        if pandas.isnull(df.Family_history_disorder_4[0]) == False:
            CR.add_paragraph(u'Autre antécédents familiaux: ' + str(df.Family_history_disorder_4[0]) + u" (degré de certitude: " + \
            str(df.Family_history_confidence_4[0]) + u"/3), chez un membre au " + str(df.Family_history_degree_4[0]) + u' degré coté ' +\
            str(df.Family_history_side_4[0]) + u'.')           
        
        
        CR.add_heading(u"Résultats de l'examen", level=1)
        db_origin = pandas.read_csv(os.getcwd().split('\\Files')[0] + '/Data/Database.csv',header = 0,engine='python',sep=';',decimal='.')
        
        CR.add_paragraph(u'Les analyses suivantes ont été effectuées en comparant le patient par rapport à une population aux caractéristiques suivantes: n = '\
        + str(len(db_origin.Age)) + '; Age moyen = ' + str("{0:.2f}".format(np.average(db.Age))) + '; ' + str("{0:.2f}".format(float(len(db_origin.Sex[db_origin.Sex == 'f']))/float(len(db.Sex))*100)) + ' % de femmes')
        

        CR.add_heading(u"Age", level=2) 
        CR.add_picture('./Files/Docs/Temporary/Age.png', width=Inches(5))      


        
        CR.add_heading(u"Personalité", level=2) 
        
        
        
        
        CR.save(os.getcwd().split('\\Files')[0] + '/Data/'+subject_code+'/'+subject_code+'_CR.docx')

finally:
    pygame.quit()
